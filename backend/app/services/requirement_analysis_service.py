"""Requirement analysis service: extract text from documents, generate test cases via DeepSeek, refine for AutoGLM."""

from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from openpyxl import load_workbook
from sqlalchemy.orm import Session

from app.config import settings
from app.models import ImportedTestCase, TestCaseFolder, TestPlanProject
from app.services.model_provider_service import ModelProviderConfig, ModelProviderService
from app.services.test_plan_import_service import TestPlanImportService


class RequirementAnalysisError(ValueError):
    pass


@dataclass
class GeneratedTestCase:
    system_name: str | None = None
    module: str | None = None
    case_name: str = ""
    precondition: str | None = None
    steps: list[str] = field(default_factory=list)
    expected_result: str = ""
    requirement_id: str | None = None
    case_type: str | None = None
    priority: str | None = None
    target_app: str | None = None
    test_module: str | None = None


@dataclass
class RequirementAnalysisResult:
    plan: TestPlanProject
    generated_count: int = 0
    raw_text_length: int = 0
    initial_case_count: int = 0
    refined_case_count: int = 0


class RequirementAnalysisService:
    """Analyze product requirement documents and generate test cases via DeepSeek."""

    FIELD_MAX_LENGTH = TestPlanImportService.FIELD_MAX_LENGTH
    MAX_RAW_TEXT_LENGTH = 30000
    MIN_RAW_TEXT_LENGTH = 50

    def __init__(self, db: Session) -> None:
        self.db = db
        self._provider_service = ModelProviderService()

    # ── Public API ──────────────────────────────────────────────────────────

    def analyze_requirement(
        self,
        file_path: Path,
        project_name: str,
        target_app: str,
        source_filename: str | None = None,
    ) -> RequirementAnalysisResult:
        """Full pipeline: extract → generate → refine → validate → save."""
        # 1. Extract text from document
        raw_text = self._extract_text(file_path)

        # 2. Generate initial test cases via DeepSeek
        initial_cases = self._generate_cases(raw_text, target_app)
        initial_count = len(initial_cases)

        # 3. Refine test cases for mobile AutoGLM execution
        refined_cases = self._refine_cases(initial_cases, target_app)
        refined_count = len(refined_cases)

        # 4. Validate and truncate fields
        validated = self._validate_and_truncate(refined_cases, target_app)
        if not validated:
            raise RequirementAnalysisError("AI 生成的测试用例经校验后无有效条目")

        # 5. Save to database
        plan = self._save_to_database(validated, project_name, target_app, source_filename)

        return RequirementAnalysisResult(
            plan=plan,
            generated_count=len(validated),
            raw_text_length=len(raw_text),
            initial_case_count=initial_count,
            refined_case_count=refined_count,
        )

    # ── Text Extraction ─────────────────────────────────────────────────────

    def _extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".txt":
            return self._extract_txt(file_path)
        if suffix == ".pdf":
            return self._extract_pdf(file_path)
        if suffix == ".docx":
            return self._extract_docx(file_path)
        if suffix in {".xlsx", ".xlsm"}:
            return self._extract_excel(file_path)
        if suffix == ".csv":
            return self._extract_csv(file_path)
        if suffix == ".doc":
            raise RequirementAnalysisError("暂不支持旧版 .doc，请另存为 .docx 后上传")
        raise RequirementAnalysisError(f"不支持的文档格式: {suffix}")

    @staticmethod
    def _extract_txt(path: Path) -> str:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="gbk", errors="replace")
        return text

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise RequirementAnalysisError("PDF 解析依赖未安装，请执行 pip install PyPDF2") from exc
        reader = PdfReader(str(path))
        pages_text: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(text)
        return "\n\n".join(pages_text)

    @staticmethod
    def _extract_docx(path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RequirementAnalysisError("Word 解析依赖未安装，请执行 pip install python-docx") from exc
        doc = Document(str(path))
        parts: list[str] = []
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    @staticmethod
    def _extract_excel(path: Path) -> str:
        workbook = load_workbook(path, read_only=True, data_only=True)
        parts: list[str] = []
        try:
            for ws in workbook.worksheets:
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(v).strip() for v in row if v is not None and str(v).strip())
                    if row_text:
                        parts.append(row_text)
        finally:
            workbook.close()
        return "\n".join(parts)

    @staticmethod
    def _extract_csv(path: Path) -> str:
        parts: list[str] = []
        with path.open("r", encoding="utf-8-sig", newline="", errors="replace") as f:
            for row in csv.reader(f):
                row_text = " | ".join(v.strip() for v in row if v.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _validate_raw_text(self, raw_text: str) -> str:
        """Validate and optionally truncate raw text."""
        text = raw_text.strip()
        if len(text) < self.MIN_RAW_TEXT_LENGTH:
            raise RequirementAnalysisError(f"需求文档内容过短（{len(text)}字符），至少需要{self.MIN_RAW_TEXT_LENGTH}字符")
        if len(text) > self.MAX_RAW_TEXT_LENGTH:
            text = text[: self.MAX_RAW_TEXT_LENGTH] + "\n\n[文档内容过长，已截取前30000字符]"
        return text

    # ── DeepSeek Calls ──────────────────────────────────────────────────────

    def _generate_cases(self, raw_text: str, target_app: str) -> list[GeneratedTestCase]:
        """Call DeepSeek to generate initial test cases from requirement text."""
        validated_text = self._validate_raw_text(raw_text)

        system_prompt = (
            "你是资深QA工程师。根据产品需求文档，生成完整的测试用例列表。\n"
            "每个测试用例必须包含以下字段：\n"
            "- case_name: 用例名称（简洁明了）\n"
            "- module: 所属功能模块\n"
            "- precondition: 前置条件（如需要登录、特定数据准备等）\n"
            "- steps: 操作步骤数组，每个元素是一个具体的操作步骤字符串\n"
            "- expected_result: 预期结果\n"
            "- priority: 优先级（高/中/低）\n"
            "- case_type: 用例类型（功能测试/界面测试/兼容性测试/性能测试）\n\n"
            "请以JSON数组格式输出，每个元素是一个测试用例对象。steps字段必须是字符串数组。\n"
            "确保覆盖所有核心功能点、边界条件和异常场景。"
        )

        user_prompt = f"目标应用：{target_app}\n\n产品需求文档内容：\n{validated_text}"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        response_text = self._call_deepseek(
            messages,
            max_tokens=settings.deepseek_req_max_tokens,
            temperature=settings.deepseek_req_temperature,
        )

        return self._parse_cases_json(response_text)

    def _refine_cases(self, cases: list[GeneratedTestCase], target_app: str) -> list[GeneratedTestCase]:
        """Call DeepSeek to refine test cases for mobile AutoGLM execution."""
        if not cases:
            return []

        cases_json = json.dumps(
            [self._case_to_dict(c) for c in cases],
            ensure_ascii=False,
            indent=2,
        )

        system_prompt = (
            "你是移动端自动化测试专家。以下是由AI生成的测试用例，请对每个测试用例进行优化改写，"
            "使其适合在移动设备上通过AutoGLM自动执行。\n\n"
            "改写要求：\n"
            "1. 每个步骤必须是明确的移动端交互操作（如：点击某按钮、在某输入框输入文字、滑动到某位置、返回上一页）\n"
            "2. 步骤描述要具体到界面元素位置（如：点击底部导航栏的「消息」标签）\n"
            "3. expected_result 要描述可观察的界面状态变化（如：页面跳转到搜索结果列表，列表显示至少一条结果）\n"
            "4. 保持原有的case_name含义，但可以补充更具体的描述\n"
            "5. 以JSON数组格式输出，字段结构同输入\n"
            "6. 不要删减用例，可以合并过于相似的用例，但总数不应少于输入的80%"
        )

        user_prompt = f"目标应用：{target_app}\n\n待改写的测试用例：\n{cases_json}"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        response_text = self._call_deepseek(
            messages,
            max_tokens=settings.deepseek_req_max_tokens,
            temperature=max(0.1, settings.deepseek_req_temperature - 0.1),  # slightly lower for refinement
        )

        return self._parse_cases_json(response_text)

    def _call_deepseek(
        self,
        messages: list[dict[str, str]],
        max_tokens: int = 4000,
        temperature: float = 0.3,
    ) -> str:
        """Call DeepSeek API via the existing ModelProviderService infrastructure."""
        if not settings.deepseek_req_enabled:
            raise RequirementAnalysisError("DeepSeek 需求分析功能未启用")
        if not settings.deepseek_req_api_key or settings.deepseek_req_api_key == "EMPTY":
            raise RequirementAnalysisError("DeepSeek API Key 未配置，请在 .env 中设置 DEEPSEEK_REQ_API_KEY")

        config = ModelProviderConfig(
            enabled=True,
            provider="openai_compatible",
            base_url=settings.deepseek_req_base_url,
            model=settings.deepseek_req_model,
            api_key=settings.deepseek_req_api_key,
            timeout_seconds=settings.deepseek_req_timeout_seconds,
            temperature=temperature,
            max_tokens=max_tokens,
        )

        client = self._provider_service.create_client(config)
        try:
            return client.complete(
                messages=messages,
                model=config.model,
                timeout_seconds=config.timeout_seconds,
                temperature=config.temperature,
                max_tokens=config.max_tokens,
            )
        except Exception as exc:
            raise RequirementAnalysisError(f"DeepSeek API 调用失败: {exc}") from exc

    # ── JSON Parsing ────────────────────────────────────────────────────────

    @staticmethod
    def _parse_cases_json(content: str) -> list[GeneratedTestCase]:
        """Parse AI response JSON into GeneratedTestCase list."""
        # Try direct JSON parse first
        payload = RequirementAnalysisService._extract_json(content)
        if payload is None:
            raise RequirementAnalysisError("AI 返回的内容无法解析为 JSON")

        # Normalize to list
        if isinstance(payload, dict):
            # Maybe wrapped in an object with a "cases" key
            if "cases" in payload and isinstance(payload["cases"], list):
                items = payload["cases"]
            elif "test_cases" in payload and isinstance(payload["test_cases"], list):
                items = payload["test_cases"]
            else:
                items = [payload]
        elif isinstance(payload, list):
            items = payload
        else:
            raise RequirementAnalysisError("AI 返回的 JSON 格式不正确，期望数组")

        cases: list[GeneratedTestCase] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            case = RequirementAnalysisService._dict_to_case(item)
            if case.case_name and case.steps and case.expected_result:
                cases.append(case)

        if not cases:
            raise RequirementAnalysisError("AI 未生成任何有效的测试用例")

        return cases

    @staticmethod
    def _extract_json(content: str) -> Any:
        """Extract JSON from AI response, handling markdown code blocks."""
        # Try markdown code block first
        match = re.search(r"```(?:json)?\s*\n?(.*?)\n?\s*```", content, flags=re.S)
        if match:
            try:
                return json.loads(match.group(1))
            except json.JSONDecodeError:
                pass

        # Try direct parse
        try:
            return json.loads(content)
        except json.JSONDecodeError:
            pass

        # Try to find JSON array
        match = re.search(r"\[.*\]", content, flags=re.S)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass

        # Try to find JSON object
        match = re.search(r"\{.*\}", content, flags=re.S)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass

        return None

    @staticmethod
    def _dict_to_case(item: dict[str, Any]) -> GeneratedTestCase:
        """Convert a dict from AI response to GeneratedTestCase."""
        steps_raw = item.get("steps", [])
        if isinstance(steps_raw, str):
            # Split by newlines or numbered items
            steps = [s.strip() for s in re.split(r"(?:\r?\n)+|(?:^|\s)(?=\d+[.、])", steps_raw) if s.strip()]
            steps = [re.sub(r"^\s*\d+[.、]\s*", "", s).strip() for s in steps if s.strip()]
        elif isinstance(steps_raw, list):
            steps = [str(s).strip() for s in steps_raw if str(s).strip()]
        else:
            steps = []

        return GeneratedTestCase(
            system_name=item.get("system_name") or item.get("module"),
            module=item.get("module"),
            case_name=str(item.get("case_name") or item.get("name") or "").strip(),
            precondition=item.get("precondition"),
            steps=steps,
            expected_result=str(item.get("expected_result") or item.get("expected") or "").strip(),
            requirement_id=item.get("requirement_id"),
            case_type=item.get("case_type") or item.get("type"),
            priority=item.get("priority"),
            target_app=item.get("target_app"),
            test_module=item.get("test_module"),
        )

    @staticmethod
    def _case_to_dict(case: GeneratedTestCase) -> dict[str, Any]:
        """Convert GeneratedTestCase to dict for JSON serialization."""
        return {
            "case_name": case.case_name,
            "module": case.module,
            "precondition": case.precondition,
            "steps": case.steps,
            "expected_result": case.expected_result,
            "priority": case.priority,
            "case_type": case.case_type,
            "requirement_id": case.requirement_id,
            "target_app": case.target_app,
            "test_module": case.test_module,
            "system_name": case.system_name,
        }

    # ── Validation ──────────────────────────────────────────────────────────

    def _validate_and_truncate(
        self, cases: list[GeneratedTestCase], target_app: str
    ) -> list[GeneratedTestCase]:
        """Validate and truncate fields to match database constraints."""
        validated: list[GeneratedTestCase] = []
        for case in cases:
            # Ensure required fields
            if not case.case_name or not case.steps or not case.expected_result:
                continue

            # Truncate fields
            case.case_name = self._truncate(case.case_name, "case_name")
            case.system_name = self._truncate(case.system_name, "system_name")
            case.module = self._truncate(case.module, "module")
            case.precondition = self._truncate(case.precondition, "precondition")
            case.expected_result = self._truncate(case.expected_result, "expected_result")
            case.requirement_id = self._truncate(case.requirement_id, "requirement_id")
            case.case_type = self._truncate(case.case_type, "case_type")
            case.priority = self._truncate(case.priority, "priority")
            case.target_app = self._truncate(case.target_app or target_app, "target_app")
            case.test_module = self._truncate(case.test_module, "test_module")

            # Truncate steps
            max_step_len = TestPlanImportService.MAX_STEP_LENGTH
            case.steps = [s[:max_step_len] for s in case.steps[: TestPlanImportService.MAX_STEPS]]

            validated.append(case)

        return validated

    def _truncate(self, value: str | None, field_name: str) -> str | None:
        """Truncate a field value to its maximum allowed length."""
        if not value:
            return None
        value = str(value).strip()
        if not value:
            return None
        max_length = self.FIELD_MAX_LENGTH.get(field_name)
        if max_length and len(value) > max_length:
            return value[: max_length - 3] + "..."
        return value

    # ── Database Save ───────────────────────────────────────────────────────

    def _save_to_database(
        self,
        cases: list[GeneratedTestCase],
        project_name: str,
        target_app: str,
        source_filename: str | None = None,
    ) -> TestPlanProject:
        """Create TestPlanProject, TestCaseFolder, and ImportedTestCase records."""
        plan = TestPlanProject(
            name=project_name.strip(),
            source_filename=source_filename,
            total_cases=0,
        )
        self.db.add(plan)
        self.db.flush()

        # Auto-create folders by grouping on module/system_name
        folder_by_group: dict[str, TestCaseFolder] = {}
        folder_sequence = 0

        imported_cases: list[ImportedTestCase] = []
        for i, case in enumerate(cases, start=1):
            # Group by module (or system_name if module is empty)
            group_key = (case.module or case.system_name or "").strip()
            if group_key and group_key not in folder_by_group:
                folder_sequence += 1
                folder = TestCaseFolder(
                    plan_id=plan.id,
                    name=group_key,
                    requirement_summary=None,
                    source_type="ai_generated",
                    source_filename=source_filename,
                    sequence=folder_sequence,
                    total_cases=0,
                )
                self.db.add(folder)
                self.db.flush()
                folder_by_group[group_key] = folder

            imported = ImportedTestCase(
                plan_id=plan.id,
                folder_id=folder_by_group[group_key].id if group_key else None,
                sequence=i,
                system_name=case.system_name,
                module=case.module,
                case_name=case.case_name,
                precondition=case.precondition,
                steps=case.steps,
                expected_result=case.expected_result,
                requirement_id=case.requirement_id,
                case_type=case.case_type,
                priority=case.priority,
                target_app=case.target_app or target_app,
                test_module=case.test_module,
                run_count=0,
                latest_result="pending",
                latest_result_note="",
            )
            self.db.add(imported)
            imported_cases.append(imported)

        if not imported_cases:
            self.db.rollback()
            raise RequirementAnalysisError("没有有效的测试用例可保存")

        plan.total_cases = len(imported_cases)

        # Update folder total_cases counts
        folder_case_counts: dict[int, int] = {}
        for ic in imported_cases:
            if ic.folder_id:
                folder_case_counts[ic.folder_id] = folder_case_counts.get(ic.folder_id, 0) + 1
        for fid, count in folder_case_counts.items():
            group_key = next((k for k, v in folder_by_group.items() if v.id == fid), None)
            if group_key and group_key in folder_by_group:
                folder_by_group[group_key].total_cases = count

        self.db.commit()
        self.db.refresh(plan)

        return plan
